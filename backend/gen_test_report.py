"""add-mvp-core 수동 검증 결과를 docx 리포트로 생성한다. 1회성 스크립트."""

import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs", "test-screenshots")
OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "TaskFlow_MVP_테스트결과.docx")

doc = Document()

title = doc.add_heading("TaskFlow MVP — 수동 검증 결과", level=0)

meta = doc.add_paragraph()
meta.add_run("대상 change: ").bold = True
meta.add_run("add-mvp-core (openspec/changes/add-mvp-core)\n")
meta.add_run("검증 환경: ").bold = True
meta.add_run("로컬 FastAPI + SQLite (uvicorn, http://127.0.0.1:8010), Chromium 기반 브라우저 자동화(Playwright)\n")
meta.add_run("검증 방식: ").bold = True
meta.add_run("백엔드 API는 curl로 직접 호출, 프론트엔드는 실제 브라우저로 화면 조작")

doc.add_paragraph()


def add_section(heading, rows, screenshots=None):
    doc.add_heading(heading, level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "검증 항목"
    hdr[1].text = "방법"
    hdr[2].text = "결과"
    for item, method, result in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = method
        cells[2].text = result
    doc.add_paragraph()
    if screenshots:
        for caption, filename in screenshots:
            path = os.path.join(SCREENSHOT_DIR, filename)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(5.5))
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)
    doc.add_paragraph()


add_section(
    "1. 인증 (user-auth)",
    [
        ("회원가입 성공 → 201 + JWT", "curl", "PASS"),
        ("이메일 중복 → 409 EMAIL_TAKEN", "curl", "PASS"),
        ("비밀번호 8자 미만 → 400 VALIDATION_ERROR", "curl", "PASS"),
        ("로그인 자격증명 오류 → 401 INVALID_CREDENTIALS", "curl + 브라우저", "PASS"),
        ("로그인 성공 → 팀 유무에 따라 team-select/app 분기", "브라우저", "PASS"),
    ],
    screenshots=[("로그인 실패 시 에러 메시지 표시", "08-login-invalid-credentials.png")],
)

add_section(
    "2. 팀 관리 (team-management)",
    [
        ("팀 생성 → invite_code 자동 발급", "curl + 브라우저", "PASS"),
        ("존재하지 않는 초대코드 → 404 NOT_FOUND", "curl", "PASS"),
        ("초대코드로 합류 → member_count 증가", "curl", "PASS"),
        ("멤버 목록에 owner 뱃지 표시", "브라우저", "PASS"),
        ("owner가 팀 탈퇴 시도 → 403 OWNER_CANNOT_LEAVE (UI에서 버튼 비활성화)", "curl + 브라우저", "PASS"),
        ("비멤버가 다른 팀 리소스 접근 → 403 FORBIDDEN", "curl", "PASS"),
    ],
    screenshots=[
        ("멤버 화면: owner는 '팀 탈퇴' 버튼이 비활성화됨", "04-members-owner-leave-disabled.png"),
        ("비멤버 접근 시 403 안내 화면", "07-forbidden-403-screen.png"),
    ],
)

add_section(
    "3. 칸반 (kanban-board)",
    [
        ("태스크 생성 (제목 + 담당자) → 201, TODO 컬럼에 표시", "브라우저", "PASS"),
        ("카드 클릭 → 상세/수정 모달 오픈", "브라우저", "PASS"),
        ("모달에서 상태 변경 + 담당자 지정 후 저장 → 컬럼 이동 확인", "브라우저", "PASS"),
        ("드래그앤드롭으로 DOING → DONE 이동", "브라우저", "PASS"),
        ("생성자/타인 삭제 권한 분기 (creator/owner만 가능, 그 외 403)", "curl", "PASS"),
        ("한글 제목 정상 저장/표시 (인코딩 문제 없음)", "curl", "PASS"),
    ],
    screenshots=[
        ("칸반: 상세 모달에서 상태·담당자 변경 후 DOING 컬럼으로 이동", "01-kanban-after-status-change.png"),
        ("칸반: 드래그앤드롭으로 DONE 컬럼 이동", "02-kanban-drag-to-done.png"),
    ],
)

add_section(
    "4. 채팅 (team-chat)",
    [
        ("메시지 전송 → 201, 화면에 말풍선으로 렌더링", "브라우저", "PASS"),
        ("1000자 초과 메시지 → 400 TOO_LONG(limit/actual 포함)", "curl", "PASS"),
        ("본인 아닌 메시지 삭제 시도 → 403 NOT_OWNER", "curl", "PASS"),
        ("본인 메시지 삭제 → 200, 화면에서 즉시 제거", "curl", "PASS"),
        ("since= 파라미터 기반 증분 폴링(5초 고정, 동적 조정 없음)", "코드 리뷰 + curl", "PASS"),
    ],
    screenshots=[("채팅: 메시지 전송 후 즉시 렌더링 (본인 메시지 삭제 버튼 노출)", "03-chat-message-sent.png")],
)

add_section(
    "5. 모바일 반응형",
    [
        ("390x844 뷰포트에서 데스크탑 네비게이션 숨김 + 햄버거 버튼 노출", "브라우저", "PASS"),
        ("햄버거 클릭 → 드로어에 칸반/채팅/멤버/로그아웃 노출", "브라우저", "PASS"),
        ("모바일에서도 칸반 3컬럼이 세로 스택으로 정상 표시", "브라우저", "PASS"),
    ],
    screenshots=[
        ("모바일(390px) 칸반 화면", "05-mobile-kanban.png"),
        ("모바일 햄버거 드로어 오픈", "06-mobile-hamburger-drawer.png"),
    ],
)

add_section(
    "6. Vercel 배포 + 프로덕션 스모크 테스트",
    [
        ("vercel integration add neon → Neon Postgres 자동 프로비저닝 + DATABASE_URL 자동 주입", "CLI", "PASS"),
        ("vercel deploy --prod → FE(정적) + BE(Python Serverless Function) 배포", "CLI", "PASS"),
        ("GET /health → 200 (SQLite 기본값일 땐 500이었으나 Neon 연결 후 정상화)", "curl", "PASS"),
        ("프로덕션에서 회원가입 → 팀 생성 → 칸반 태스크 생성(한글) → 채팅 전송", "브라우저(Playwright)", "PASS"),
        ("프로덕션에서 owner 탈퇴 시도 → 403 OWNER_CANNOT_LEAVE", "curl", "PASS"),
        ("프로덕션에서 비멤버 접근 → 403 FORBIDDEN", "curl", "PASS"),
    ],
    screenshots=[("배포된 Vercel URL(https://taskflow-openspec-ochre.vercel.app)에서 채팅 전송 확인", "09-vercel-prod-chat.png")],
)

add_section(
    "7. 루트 경로 404 수정 + 프로덕션 모바일 재확인",
    [
        ("배포 도메인 루트(/) 접근 시 404 발생 → frontend/index.html 추가로 수정", "curl + 브라우저", "FIXED"),
        ("비로그인 상태로 루트(/) 접근 → /login.html로 리다이렉트", "브라우저", "PASS"),
        ("로그인 상태로 루트(/) 접근 → /app.html로 리다이렉트", "브라우저", "PASS"),
        ("프로덕션 URL, 390x844 모바일 뷰포트에서 로그인 화면 정상 렌더링", "브라우저", "PASS"),
        ("프로덕션 URL, 모바일 칸반 화면 (햄버거 버튼, 세로 스택 컬럼)", "브라우저", "PASS"),
        ("프로덕션 URL, 모바일 햄버거 드로어 오픈", "브라우저", "PASS"),
        ("프로덕션 URL, 모바일 드로어에서 채팅 화면으로 이동", "브라우저", "PASS"),
    ],
    screenshots=[
        ("프로덕션 모바일: 로그인 화면", "10-mobile-prod-login.png"),
        ("프로덕션 모바일: 칸반 화면", "11-mobile-prod-kanban.png"),
        ("프로덕션 모바일: 햄버거 드로어", "12-mobile-prod-drawer.png"),
        ("프로덕션 모바일: 채팅 화면", "13-mobile-prod-chat.png"),
    ],
)

add_section(
    "8. 로컬 재검증 (배포 관련 코드 변경 이후)",
    [
        ("database.py/main.py의 Vercel 대응 변경(postgres:// 정규화, VERCEL 조건부 StaticFiles) 이후 로컬 서버 재기동", "uvicorn", "PASS"),
        ("회원가입 성공/이메일 중복(409)/약한 비밀번호(400)", "curl", "PASS"),
        ("로그인 성공/실패(401 INVALID_CREDENTIALS)", "curl", "PASS"),
        ("팀 생성 → 초대코드 합류(성공/404) → 멤버 목록 조회", "curl", "PASS"),
        ("태스크 생성(assignee) → 필터(@me/미할당) → 상태변경(PATCH) → 삭제 권한(403→200)", "curl", "PASS"),
        ("채팅 전송 → 1000자 초과(400 TOO_LONG) → 삭제 권한(403 NOT_OWNER→200)", "curl", "PASS"),
        ("owner 탈퇴 차단(403) / 비멤버 팀 접근 차단(403)", "curl", "PASS"),
        ("root(/) → /login.html 리다이렉트, 회원가입 → 팀 생성 → 칸반 → 채팅 → 멤버 화면 전체 플로우", "브라우저(Playwright)", "PASS"),
        ("브라우저 콘솔 에러 0건 (favicon 404 제외)", "브라우저", "PASS"),
    ],
    screenshots=[("로컬 재검증: 멤버 화면 (owner 팀 탈퇴 버튼 비활성화)", "14-local-recheck-members.png")],
)

doc.add_heading("종합", level=1)
summary = doc.add_paragraph()
summary.add_run(
    "백엔드 API 18개(Auth 4 + Team 5 + Task 6 + Chat 3)와 프론트엔드 9개 화면 모두 spec-driven 스펙에 정의된 "
    "정상/에러 시나리오대로 동작함을 확인했다. 결정 #9(owner 탈퇴 차단), 결정 #10(5초 고정 폴링) 모두 코드와 "
    "화면에 그대로 반영되어 있다. Vercel(FE+BE) + Neon(Postgres) 배포까지 완료했고, 배포된 URL에서 동일한 "
    "시나리오를 브라우저 자동화와 curl로 재검증했다.\n"
)
summary.add_run("배포 URL: ").bold = True
summary.add_run("https://taskflow-openspec-ochre.vercel.app")

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
